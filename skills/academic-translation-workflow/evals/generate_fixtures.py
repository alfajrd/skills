"""
One-off generator for synthetic Indonesian academic-article fixtures used by
the eval test cases. Not part of the shipped skill - just test data.

Produces evals/fixtures/gotong_royong.docx (headings, body prose with an
abbreviation + decimal number to stress the sentence splitter, a bullet
list, a data table, an image + caption, a bibliography section, and a
footnote) and evals/fixtures/mangrove.docx (a simpler second variant).
"""
import io
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image

FIXTURES = Path(__file__).parent / "fixtures"
FIXTURES.mkdir(exist_ok=True)


def make_placeholder_png(path, text, size=(400, 250), color=(90, 130, 170)):
    img = Image.new("RGB", size, color=color)
    img.save(path)


def build_gotong_royong():
    doc = Document()
    doc.add_heading("Gotong Royong sebagai Modal Sosial dalam Pembangunan Desa", level=0)

    doc.add_heading("Pendahuluan", level=1)
    doc.add_paragraph(
        "Gotong royong merupakan salah satu nilai budaya yang telah lama menjadi "
        "fondasi kehidupan sosial masyarakat pedesaan di Indonesia. Praktik ini "
        "tercatat telah berkontribusi terhadap 3.5 persen peningkatan partisipasi "
        "warga dalam proyek pembangunan desa, menurut Susanto dkk. Penelitian ini "
        "bertujuan untuk mengkaji bagaimana gotong royong dapat dimanfaatkan "
        "sebagai modal sosial dalam program pembangunan berkelanjutan."
    )
    doc.add_paragraph(
        "Selain gotong royong, terdapat pula istilah sambatan yang digunakan di "
        "beberapa wilayah Jawa untuk menggambarkan praktik serupa. Kedua istilah "
        "ini mencerminkan nilai kolektivitas yang kuat dalam masyarakat agraris."
    )

    doc.add_heading("Metode", level=1)
    doc.add_paragraph(
        "Penelitian ini menggunakan pendekatan kualitatif dengan studi kasus di "
        "tiga desa. Data dikumpulkan melalui wawancara mendalam dan observasi "
        "partisipatif selama enam bulan."
    )
    doc.add_paragraph("Tahapan penelitian meliputi:", style=None)
    for item in [
        "Identifikasi lokasi penelitian",
        "Pengumpulan data lapangan",
        "Analisis tematik terhadap hasil wawancara",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Desa"
    hdr[1].text = "Jumlah Partisipan"
    for desa, jumlah in [("Desa Sukamaju", "42"), ("Desa Cibunar", "37"), ("Desa Margasari", "51")]:
        row = table.add_row().cells
        row[0].text = desa
        row[1].text = jumlah

    img_path = FIXTURES / "_chart.png"
    make_placeholder_png(img_path, "chart")
    doc.add_picture(str(img_path), width=Inches(3))
    cap = doc.add_paragraph("Gambar 1. Grafik partisipasi warga per desa.")
    cap.style = doc.styles["Caption"] if "Caption" in [s.name for s in doc.styles] else cap.style

    doc.add_heading("Kesimpulan", level=1)
    doc.add_paragraph(
        "Gotong royong terbukti menjadi modal sosial yang efektif dalam "
        "mendorong partisipasi warga desa. Pemerintah daerah disarankan untuk "
        "mengintegrasikan nilai ini ke dalam kebijakan pembangunan desa."
    )

    doc.add_heading("Daftar Pustaka", level=1)
    doc.add_paragraph("Susanto, A., Wibowo, B., & Rahayu, C. (2019). Modal sosial dan pembangunan desa. Jurnal Sosiologi Pedesaan, 12(1), 45-60.")
    doc.add_paragraph("Hidayat, R. (2021). Gotong royong dalam perspektif antropologi. Yogyakarta: Gadjah Mada University Press.")

    out_path = FIXTURES / "gotong_royong.docx"
    doc.save(out_path)
    img_path.unlink(missing_ok=True)
    _inject_footnote(
        out_path,
        target_paragraph_text_startswith="Gotong royong merupakan salah satu",
        footnote_text="Istilah ini juga dikenal dengan nama “sambatan” di beberapa daerah Jawa.",
    )
    print(f"Wrote {out_path}")


def build_mangrove():
    doc = Document()
    doc.add_heading("Reboisasi Mangrove untuk Mitigasi Abrasi Pantai", level=0)

    doc.add_heading("Latar Belakang", level=1)
    doc.add_paragraph(
        "Abrasi pantai menjadi ancaman serius bagi wilayah pesisir Indonesia. "
        "Reboisasi mangrove dinilai sebagai solusi berbasis alam yang efektif "
        "dan berbiaya rendah dibandingkan pembangunan tanggul beton."
    )
    doc.add_paragraph(
        "Program reboisasi mangrove yang melibatkan masyarakat lokal, sering "
        "disebut sebagai program berbasis komunitas, telah menunjukkan hasil "
        "yang menjanjikan di beberapa provinsi."
    )

    doc.add_heading("Hasil", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Lokasi"
    hdr[1].text = "Luas (ha)"
    hdr[2].text = "Tingkat Keberhasilan"
    for lokasi, luas, keberhasilan in [
        ("Pesisir Demak", "12.5", "78%"),
        ("Pesisir Indramayu", "8.2", "65%"),
    ]:
        row = table.add_row().cells
        row[0].text = lokasi
        row[1].text = luas
        row[2].text = keberhasilan

    doc.add_heading("Daftar Pustaka", level=1)
    doc.add_paragraph("Prasetyo, D. (2020). Ekologi mangrove pesisir utara Jawa. Bandung: ITB Press.")

    out_path = FIXTURES / "mangrove.docx"
    doc.save(out_path)
    print(f"Wrote {out_path}")


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
R_NS = "http://schemas.openxmlformats.org/package/2006/relationships"

FOOTNOTES_XML_TEMPLATE = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="{w}">
  <w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>
  <w:footnote w:id="2"><w:p><w:r><w:t xml:space="preserve">{text}</w:t></w:r></w:p></w:footnote>
</w:footnotes>"""


def _inject_footnote(docx_path, target_paragraph_text_startswith, footnote_text):
    """Hand-injects a minimal but valid footnotes.xml part + a footnote
    reference run into the target paragraph, plus the content-types and
    relationship wiring Word requires. python-docx has no native footnote
    API, so this is done via direct zip/XML surgery - only needed here to
    build a realistic test fixture; real client documents already have this
    wiring if they contain footnotes."""
    from lxml import etree

    with zipfile.ZipFile(docx_path) as z:
        names = z.namelist()
        data = {n: z.read(n) for n in names}

    doc_xml = etree.fromstring(data["word/document.xml"])
    ns = {"w": W_NS}
    target_p = None
    for p in doc_xml.iter(f"{{{W_NS}}}p"):
        texts = "".join(t.text or "" for t in p.iter(f"{{{W_NS}}}t"))
        if texts.startswith(target_paragraph_text_startswith):
            target_p = p
            break
    if target_p is None:
        raise ValueError("target paragraph not found for footnote injection")

    run = etree.SubElement(target_p, f"{{{W_NS}}}r")
    rpr = etree.SubElement(run, f"{{{W_NS}}}rPr")
    etree.SubElement(rpr, f"{{{W_NS}}}vertAlign", {f"{{{W_NS}}}val": "superscript"})
    etree.SubElement(run, f"{{{W_NS}}}footnoteReference", {f"{{{W_NS}}}id": "2"})
    data["word/document.xml"] = etree.tostring(doc_xml, xml_declaration=True, encoding="UTF-8", standalone=True)

    data["word/footnotes.xml"] = FOOTNOTES_XML_TEMPLATE.format(w=W_NS, text=footnote_text).encode("utf-8")

    ct_xml = etree.fromstring(data["[Content_Types].xml"])
    override = etree.SubElement(ct_xml, f"{{{CT_NS}}}Override")
    override.set("PartName", "/word/footnotes.xml")
    override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml")
    data["[Content_Types].xml"] = etree.tostring(ct_xml, xml_declaration=True, encoding="UTF-8", standalone=True)

    rels_path = "word/_rels/document.xml.rels"
    rels_xml = etree.fromstring(data[rels_path])
    existing_ids = [r.get("Id") for r in rels_xml]
    new_id = "rIdFootnotes1"
    assert new_id not in existing_ids
    rel = etree.SubElement(rels_xml, f"{{{R_NS}}}Relationship")
    rel.set("Id", new_id)
    rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes")
    rel.set("Target", "footnotes.xml")
    data[rels_path] = etree.tostring(rels_xml, xml_declaration=True, encoding="UTF-8", standalone=True)

    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, content in data.items():
            zout.writestr(name, content)


if __name__ == "__main__":
    build_gotong_royong()
    build_mangrove()
