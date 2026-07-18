# -*- coding: utf-8 -*-
"""
Usage: python build_catatan.py <reviewed.json> <editorial_notes.json> <output_catatan.docx>

Builds the "Catatan Penyuntingan" (editing-notes) document that MUST
accompany every final deliverable (v3_final.docx). The final document has
some passages highlighted in yellow - places where the translation departed
from a plain sentence-for-sentence rendering and the author needs to look:

  - a missing citation flagged as "[butuh sitasi]" (citation needed)
  - a predicted/inferred citation the author should verify
  - a sentence rewritten because the source was confusing
  - a sentence split into several because it carried more than one main idea
    (or several merged into one)
  - anything else the translator wants the author to confirm

The highlight alone shows the author WHERE to look; this document tells them
WHY, one entry per highlighted span. Without it, a yellow highlight in the
final file is a mystery.

Inputs
------
reviewed.json      : output of read_reviewed_doc.py. Each row's `resolved`
                     text carries `{{...}}` markup for every highlighted span
                     (this is what the rebuild renders as yellow highlight in
                     the final document), so the highlights are enumerable
                     straight from here without re-parsing the .docx.
editorial_notes.json : a list Claude maintains while translating (Step 3),
                     one entry per highlight, explaining it:
  [
    {
      "id": "p24.s2",
      "category": "predicted-citation",
      "highlighted": "(Ndlovu-Gatsheni, 2020)",
      "note": "Sumber memiliki celah sitasi ('...setiap hari. . menunjukkan...'); disimpulkan dari istilah 'cognitive empire' yang cocok dengan daftar pustaka. Mohon verifikasi."
    },
    ...
  ]
  `id` matches the reviewed-row id. `highlighted` is the exact inner text of
  the `{{...}}` span this note explains (needed only to disambiguate when a
  single row has more than one highlight; may be omitted otherwise). Every
  highlight in the final document should have a matching note, and every note
  should still match a live highlight - the script warns about either kind of
  mismatch rather than guessing.

Known categories (others are allowed; the raw string is shown as-is):
  needs-citation, predicted-citation, rewritten, split, merged, ambiguity, other
"""
import sys
import json
import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor

from docx_utils import apply_translated_text, extract_highlights, strip_markup

CATEGORY_LABELS = {
    "needs-citation": "Butuh sitasi / Citation needed",
    "predicted-citation": "Prediksi sitasi – verifikasi / Predicted citation – verify",
    "rewritten": "Kalimat ditulis ulang / Sentence rewritten",
    "split": "Kalimat dipecah / Sentence split (>1 main idea)",
    "merged": "Kalimat digabung / Sentences merged",
    "ambiguity": "Makna ambigu / Ambiguous meaning",
    "other": "Lainnya / Other",
}


def _label(category):
    return CATEGORY_LABELS.get(category, category or "—")


def build_catatan(reviewed, notes, out_path):
    """Reconcile notes against the highlighted spans in reviewed rows and
    build the catatan document. Returns (entries, stale_notes, unexplained)
    so the caller can warn about mismatches."""
    rows_by_id = {r["id"]: r for r in reviewed.get("rows", [])}

    # Enumerate every highlighted span in document order: (row, span_text).
    highlights = []  # (row_id, span_text, resolved_text)
    for r in reviewed.get("rows", []):
        for span in extract_highlights(r.get("resolved", "") or ""):
            highlights.append((r["id"], span, r.get("resolved", "")))

    notes = list(notes)
    note_used = [False] * len(notes)

    def find_note(row_id, span_text):
        # Prefer a note that names this exact span; fall back to a note for
        # this row with no `highlighted` given (unambiguous single-highlight
        # rows), preferring not-yet-used notes.
        for prefer_exact in (True, False):
            for i, n in enumerate(notes):
                if note_used[i] or n.get("id") != row_id:
                    continue
                hl = n.get("highlighted")
                if prefer_exact:
                    if hl is not None and (hl == span_text or hl in span_text or span_text in hl):
                        return i
                else:
                    if hl is None:
                        return i
        return None

    entries = []  # (row_id, span_text, resolved_text, category, note_text)
    for row_id, span_text, resolved in highlights:
        i = find_note(row_id, span_text)
        if i is not None:
            note_used[i] = True
            entries.append((row_id, span_text, resolved, notes[i].get("category", "other"), notes[i].get("note", "")))
        else:
            entries.append((row_id, span_text, resolved, "other",
                            "(TIDAK ADA CATATAN / no editorial note found for this highlight - explain or remove the highlight)"))

    unexplained = [(rid, span) for (rid, span, _r, _c, note) in entries if note.startswith("(TIDAK ADA CATATAN")]
    stale_notes = [notes[i] for i, used in enumerate(note_used) if not used]

    _render(reviewed, entries, out_path)
    return entries, stale_notes, unexplained


def _render(reviewed, entries, out_path):
    doc = Document()

    doc.add_heading("Catatan Penyuntingan", level=0)
    sub = doc.add_paragraph()
    src = reviewed.get("source_file") or ""
    if src:
        sub.add_run(f"Dokumen sumber / Source file: {src}").italic = True

    doc.add_paragraph(
        "Dokumen ini menjelaskan bagian-bagian yang disorot kuning pada "
        "terjemahan akhir. Setiap sorotan adalah tempat penerjemah menyimpang "
        "dari penerjemahan kalimat-per-kalimat biasa dan meminta penulis "
        "memeriksanya – misalnya sitasi yang hilang atau diprediksi, atau "
        "kalimat yang ditulis ulang / dipecah karena mengandung lebih dari "
        "satu gagasan. Mohon periksa setiap butir dan konfirmasi atau perbaiki."
    )
    doc.add_paragraph(
        "This document explains every yellow-highlighted passage in the final "
        "translation. Each highlight marks a place where the translation "
        "departed from a plain sentence-for-sentence rendering and needs the "
        "author's check – e.g. a missing or predicted citation, or a "
        "sentence rewritten / split because it carried more than one idea. "
        "Please review each item and confirm or correct it."
    ).runs[0].italic = True

    if not entries:
        doc.add_paragraph(
            "Tidak ada bagian yang disorot pada terjemahan ini. / "
            "No highlighted passages in this translation."
        )
        doc.save(out_path)
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    widths = [Inches(0.4), Inches(2.6), Inches(1.7), Inches(2.3)]
    headers = ["No.", "Bagian yang disorot / Highlighted", "Kategori / Category", "Catatan / Note"]
    for cell, text, width in zip(table.rows[0].cells, headers, widths):
        cell.text = text
        cell.width = width
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    for n, (row_id, span_text, resolved, category, note_text) in enumerate(entries, start=1):
        cells = table.add_row().cells
        cells[0].text = str(n)

        # Highlighted span (bold), then the surrounding sentence for context.
        hl_para = cells[1].paragraphs[0]
        apply_translated_text(hl_para, f"*{span_text}*")  # show it emphasized
        ctx = strip_markup(resolved).strip()
        if ctx and ctx != strip_markup(span_text).strip():
            ctx_para = cells[1].add_paragraph()
            run = ctx_para.add_run(f"Konteks: “{ctx}”")
            run.italic = True
            run.font.size = Pt(8)

        cells[2].text = _label(category)
        apply_translated_text(cells[3].paragraphs[0], note_text)

        for cell, width in zip(cells, widths):
            cell.width = width

    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("reviewed_json")
    parser.add_argument("editorial_notes_json")
    parser.add_argument("output_docx")
    args = parser.parse_args()

    reviewed = json.loads(Path(args.reviewed_json).read_text(encoding="utf-8"))
    notes_path = Path(args.editorial_notes_json)
    notes = json.loads(notes_path.read_text(encoding="utf-8")) if notes_path.exists() else []

    entries, stale_notes, unexplained = build_catatan(reviewed, notes, args.output_docx)

    print(f"Wrote catatan penyuntingan with {len(entries)} highlighted entr(ies) to {args.output_docx}")
    if unexplained:
        print(
            f"WARNING: {len(unexplained)} highlight(s) in the final document have NO matching "
            f"editorial note - they appear in the catatan flagged as unexplained. Add a note for "
            f"each (or remove the highlight): {unexplained}"
        )
    if stale_notes:
        ids = [n.get("id") for n in stale_notes]
        print(
            f"WARNING: {len(stale_notes)} editorial note(s) did not match any live highlight in the "
            f"final text (the highlighted span may have been edited away during review) - they were "
            f"NOT included in the catatan: {ids}"
        )


if __name__ == "__main__":
    main()
