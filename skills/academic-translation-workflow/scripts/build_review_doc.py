"""
Usage: python build_review_doc.py <annotated.json> <output_review.docx>

Builds the reviewer-facing Word document: a front-matter briefing (topic
overview, key terminology, open questions) followed by the segmented
translation table (ID | Type | Source | Proposed | Final, Final left empty).

annotated.json is segments.json's rows PLUS Claude's translation work:
{
  "source_file": "...", "client": "...", "title": "...",
  "topic_overview": "one or more paragraphs, blank-line separated, *term* for italics",
  "terminology": [{"term": "...", "note": "..."}],
  "questions": [{"question": "..."}],
  "rows": [{"id":"p3.s0","type":"body","source":"...","proposed":"...","final":""}, ...]
}

Foreign-word italics + first-occurrence gloss: this script does NOT decide
what to italicize - it only renders whatever *markup* Claude already put in
`proposed` / `topic_overview` / `note` strings. Claude is responsible for:
italicizing every foreign-language word/phrase, and following the FIRST one
in the whole document with a plain-language gloss in parentheses, e.g.
"*gotong royong* (communal work)". Every later occurrence of that same term
should stay italic but NOT repeat the gloss.
"""
import sys
import json
import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx_utils import apply_translated_text


def add_markup_paragraph(doc_or_cell, text, style=None):
    if hasattr(doc_or_cell, "add_paragraph"):
        p = doc_or_cell.add_paragraph(style=style)
    else:
        p = doc_or_cell.paragraphs[0]
    apply_translated_text(p, text)
    return p


def set_cell_markup(cell, text):
    apply_translated_text(cell.paragraphs[0], text)


def build_review_doc(annotated, out_path):
    doc = Document()

    title_p = doc.add_heading(level=0)
    apply_translated_text(title_p, annotated.get("title") or "Translation Brief")
    meta = doc.add_paragraph()
    meta.add_run(f"Source file: {annotated.get('source_file', '')}").italic = True
    if annotated.get("client"):
        meta.add_run(f"   |   Client: {annotated['client']}").italic = True

    doc.add_heading("Topic Overview", level=1)
    for para in (annotated.get("topic_overview") or "").split("\n\n"):
        para = para.strip()
        if para:
            add_markup_paragraph(doc, para)

    terminology = annotated.get("terminology") or []
    if terminology:
        doc.add_heading("Key Terminology", level=1)
        doc.add_paragraph(
            "Terms below either lack a clean single-word English equivalent, "
            "are used in a specific technical sense in this article, or are "
            "kept as loanwords per academic convention."
        )
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Term"
        hdr[1].text = "Notes / English equivalent"
        for item in terminology:
            row = table.add_row().cells
            set_cell_markup(row[0], item.get("term", ""))
            set_cell_markup(row[1], item.get("note", ""))

    questions = annotated.get("questions") or []
    doc.add_heading("Questions for Translator", level=1)
    if questions:
        doc.add_paragraph(
            "Please answer directly in the table below (right-hand column), "
            "then resend this document."
        )
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Question"
        hdr[1].text = "Your Answer"
        for q in questions:
            row = table.add_row().cells
            set_cell_markup(row[0], q.get("question", ""))
            row[1].text = ""
    else:
        doc.add_paragraph("None - nothing ambiguous surfaced in this pass.")

    doc.add_page_break()
    doc.add_heading("Segmented Translation", level=1)
    doc.add_paragraph(
        "Fill in the \"Final\" column: write \"default\" to accept the "
        "Proposed translation as-is, or write your own replacement. "
        "Row IDs are stable references back to the source document - "
        "use them if you need to ask about a specific row."
    )
    if any(r.get("citation_fields") for r in (annotated.get("rows") or [])):
        doc.add_paragraph(
            "Rows marked \"[citation]\" in the Type column contain a live "
            "reference-manager citation (Zotero/Mendeley/EndNote). If you "
            "override the Proposed text, reproduce that citation's text "
            "exactly - rephrasing it (even just formatting) will break its "
            "live link back to your reference library in the final document."
        )

    rows = annotated.get("rows") or []
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    widths = [Inches(0.6), Inches(0.7), Inches(2.1), Inches(2.1), Inches(2.1)]
    hdr_cells = table.rows[0].cells
    for cell, text, width in zip(hdr_cells, ["ID", "Type", "Source (ID)", "Proposed (EN)", "Final (EN)"], widths):
        cell.text = text
        cell.width = width
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    for row_data in rows:
        cells = table.add_row().cells
        cells[0].text = row_data.get("id", "")
        cells[1].text = row_data.get("type", "") + (" [citation]" if row_data.get("citation_fields") else "")
        cells[2].text = row_data.get("source", "")
        set_cell_markup(cells[3], row_data.get("proposed", ""))
        cells[4].text = row_data.get("final", "")
        for cell, width in zip(cells, widths):
            cell.width = width

    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("annotated_json")
    parser.add_argument("output_docx")
    args = parser.parse_args()

    annotated = json.loads(Path(args.annotated_json).read_text(encoding="utf-8"))
    build_review_doc(annotated, args.output_docx)
    print(f"Wrote review document to {args.output_docx}")


if __name__ == "__main__":
    main()
