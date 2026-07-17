"""
Shared helpers for the academic-translation-workflow skill.

Design notes (read this before modifying):
- Paragraphs are visited in true document order (body paragraphs AND table
  cells, interleaved correctly) via `iter_block_items`. Each visited unit gets
  a stable string ID derived purely from document structure/position, so the
  SAME ids can be re-derived by re-opening the ORIGINAL, untouched source file
  later during reconstruction. Do not switch to python-docx's separate
  `.paragraphs` / `.tables` lists elsewhere in this skill - they don't
  preserve interleaving order and would desync the ids.
- Footnotes live in a separate OOXML part (word/footnotes.xml) that
  python-docx does not model, so they're handled by direct XML access via
  zipfile + lxml, not through the Document object.
- Italics are expressed with a lightweight `*term*` markup (like Markdown)
  inside any "proposed" / "final" string, rather than character offsets.
  This is far easier for an LLM to produce reliably. `runs_from_markup`
  turns that markup into actual italic runs.
"""
import re
import copy
import zipfile
import shutil
from lxml import etree

from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_COLOR_INDEX

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# --- Document-order traversal -------------------------------------------------

def _iter_block_items(parent_elm, parent):
    for child in parent_elm.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):
            yield Table(child, parent)


def iter_units(document):
    """Yield (id, kind, paragraph) for every translatable unit in document order.

    kind is one of: 'body_para', 'table_cell_para'.
    id looks like 'p3' for body paragraph #3 (0-indexed over body paragraphs
    only), or 'tbl0_r1_c2_p0' for the 1st paragraph of the cell at row1/col2
    of the 1st top-level table. One level of table nesting only (nested
    tables inside a cell are skipped - known limitation, flag for manual
    review if encountered).
    """
    p_idx = 0
    t_idx = 0
    for item in _iter_block_items(document.element.body, document):
        if isinstance(item, Paragraph):
            yield f"p{p_idx}", "body_para", item
            p_idx += 1
        elif isinstance(item, Table):
            for r_idx, row in enumerate(item.rows):
                for c_idx, cell in enumerate(row.cells):
                    for cp_idx, cell_para in enumerate(cell.paragraphs):
                        yield f"tbl{t_idx}_r{r_idx}_c{c_idx}_p{cp_idx}", "table_cell_para", cell_para
            t_idx += 1


# --- Paragraph classification --------------------------------------------------

BIBLIOGRAPHY_HEADINGS = re.compile(
    r'daftar\s+pustaka|bibliograf|^references$', re.IGNORECASE
)
# Matches a paragraph whose ENTIRE text is just a section-marker word, e.g. a
# bare "REFERENCES" paragraph styled as plain Normal/Body Text rather than a
# real Word heading - common in draft manuscripts that don't use built-in
# heading styles consistently. Anchored on the full string (not .search) so
# it can't misfire on "...see References above" appearing mid-sentence.
BIBLIOGRAPHY_MARKER_ONLY = re.compile(
    r'^(daftar\s+pustaka|bibliograf\w*|references?)$', re.IGNORECASE
)
LIST_TEXT_RE = re.compile(r'^([•●\-\*]|\d+[.)])\s+')


def classify_paragraph(paragraph, in_bibliography):
    """Return one of: empty, heading, heading_starts_bibliography, caption,
    quote, list_item, bibliography, body."""
    text = paragraph.text.strip()
    if not text:
        return "empty"
    style = (paragraph.style.name or "").lower() if paragraph.style else ""
    if "heading" in style or "title" in style:
        if BIBLIOGRAPHY_HEADINGS.search(text):
            return "heading_starts_bibliography"
        return "heading"
    if BIBLIOGRAPHY_MARKER_ONLY.match(text):
        return "heading_starts_bibliography"
    if "caption" in style:
        return "caption"
    if "quote" in style:
        return "quote"
    if "list" in style or LIST_TEXT_RE.match(text):
        return "list_item"
    if in_bibliography:
        return "bibliography"
    return "body"


# --- Sentence splitting (heuristic - Claude should sanity-check output) -------

_ABBREVIATIONS = {
    "dr", "prof", "ir", "drs", "dra", "h", "hj", "sh", "se", "si", "mm", "msi",
    "ma", "md", "dkk", "dll", "dst", "yth", "no", "jl", "st", "al", "vol",
    "hal", "cf", "vs", "tbk", "pt", "cv",
}

_SENTENCE_BOUNDARY_RE = re.compile(
    r'(?<=[.!?])\s+(?=[A-Z0-9"“‘(])'
)


def split_sentences(text):
    """Heuristic sentence splitter tuned for Indonesian academic prose.
    Not perfect around abbreviations/decimals - segment_docx.py's output
    (segments.json) should be reviewed and fixed up (merge stray splits)
    before translation proceeds, rather than trusted blindly.
    """
    text = text.strip()
    if not text:
        return []
    protected = re.sub(r'(?<=\d)\.(?=\d)', '', text)  # protect decimals
    parts = _SENTENCE_BOUNDARY_RE.split(protected)

    sentences = []
    buf = ""
    for part in parts:
        buf = f"{buf} {part}".strip() if buf else part
        m = re.search(r'\b([A-Za-z]+)\.\s*$', buf)
        if m and m.group(1).lower() in _ABBREVIATIONS:
            continue  # likely an abbreviation, not a real sentence end - keep accumulating
        sentences.append(buf)
        buf = ""
    if buf:
        sentences.append(buf)
    return [s.replace('', '.').strip() for s in sentences if s.strip()]


# --- Reference-manager in-text citations (Zotero / Mendeley / EndNote) -------
#
# These tools insert in-text citations as OOXML *fields*, not plain text: a
# `w:fldChar type="begin"` run, one or more `w:instrText` runs carrying the
# actual instruction (e.g. `ADDIN ZOTERO_ITEM CSL_CITATION {...json...}` or
# `ADDIN EN.CITE` for Mendeley/EndNote), a `w:fldChar type="separate"` run,
# then the visible "result" text (e.g. "(Bhambra, 2021)") in ordinary `w:t`
# runs, and finally `w:fldChar type="end"`. That structure is what keeps the
# citation "live" - linked back to the user's reference library so they can
# still update/restyle it later. python-docx's Run.text already ignores
# w:instrText (only w:t/w:tab/w:br/w:cr/w:noBreakHyphen/w:ptab count), so
# *reading* a paragraph's text already surfaces just the visible citation
# text correctly - no special handling needed there. The risk is on
# *rebuild*: naively stripping all runs and replacing them with plain
# translated text would silently flatten every citation field to dead text.
# Mendeley/EndNote sometimes nest a second, hidden field inside the visible
# one; tracking begin/end at matching depth (not just the first "end" seen)
# treats that whole nested structure as one atomic span, which is exactly
# what we want - we never need to understand a field's internals, only its
# boundaries, since we're preserving it unchanged rather than parsing it.

def iter_field_spans(paragraph):
    """Yield (visible_text, run_elements) for each top-level OOXML field
    found among this paragraph's runs, in document order. `run_elements` is
    the list of raw `<w:r>` elements spanning that field (begin...end,
    inclusive) - callers splice these back in as-is rather than rebuilding
    them, since a field's internal structure is reference-manager-specific
    and not worth (or safe to) reverse-engineer.
    """
    runs = paragraph.runs
    i = 0
    while i < len(runs):
        fld = runs[i]._r.find(qn('w:fldChar'))
        if fld is not None and fld.get(qn('w:fldCharType')) == 'begin':
            depth = 1
            j = i + 1
            while j < len(runs) and depth > 0:
                fldj = runs[j]._r.find(qn('w:fldChar'))
                if fldj is not None:
                    ftype = fldj.get(qn('w:fldCharType'))
                    if ftype == 'begin':
                        depth += 1
                    elif ftype == 'end':
                        depth -= 1
                j += 1
            end = j - 1  # index of the matching (depth-0) end fldChar run
            span_runs = runs[i:end + 1]
            visible = "".join(run.text for run in span_runs).strip()
            if visible:
                yield visible, [r._r for r in span_runs]
            i = end + 1
        else:
            i += 1


# --- Writing translated text back into a paragraph, with *italic* markup ------

_MARKUP_RE = re.compile(r'\*([^*]+)\*')
_HIGHLIGHT_RE = re.compile(r'\{\{(.+?)\}\}', re.DOTALL)
# `{{text}}` renders as yellow-highlighted - used for editorial flags like
# "{{[butuh sitasi]}}" (citation needed) that an author/editor needs to
# visually notice, distinct from *text*'s italic (foreign-word) convention.
# Curly braces were chosen specifically so the highlighted content can
# itself contain literal square brackets (e.g. "[butuh sitasi]") without
# any delimiter collision. The two markers aren't designed to nest in the
# same span for this skill's use cases (highlight is split out first,
# italic within each resulting piece), so `*{{x}}*` or `{{*x*}}` will only
# apply the outer marker correctly - keep the two conventions on separate
# spans of text.


def text_with_markup(container):
    """Inverse of apply_translated_text: read a paragraph (or a table cell,
    which may hold multiple paragraphs) back out as a plain string, but
    re-encode any italic run as `*text*` and any highlighted run as
    `{{text}}` (both, nested, as `{{*text*}}`). This is what lets
    italic-ness/highlight-ness survive the round trip through Word - a
    translator who accepts a row as-is, or who types their own override and
    applies Ctrl+I or a highlight color, ends up with the same markup
    representation that apply_translated_text expects on the way back in.
    Consecutive runs sharing the same italic+highlight state are merged
    first so a span split across multiple runs (e.g. by spellcheck) doesn't
    turn into several adjacent *fragments*.
    """
    paragraphs = container.paragraphs if hasattr(container, "paragraphs") else [container]
    out = []
    for p in paragraphs:
        merged = []
        for r in p.runs:
            if not r.text:
                continue
            state = (bool(r.italic), r.font.highlight_color is not None)
            if merged and merged[-1][1] == state:
                merged[-1][0] += r.text
            else:
                merged.append([r.text, state])
        for text, (italic, highlighted) in merged:
            if italic:
                text = f"*{text}*"
            if highlighted:
                text = f"{{{{{text}}}}}"
            out.append(text)
    return "".join(out)


def apply_translated_text(paragraph, text):
    """Replace a paragraph's runs with `text`, preserving the paragraph's
    first run's character formatting (font/size/color/bold/etc.) as the base
    style, and rendering `*term*` spans as italic runs. Paragraph-level
    formatting (alignment, style, numbering, spacing) is untouched since we
    never touch paragraph-level XML, only its runs.

    Any reference-manager citation field (see iter_field_spans above) found
    in the ORIGINAL paragraph is preserved as a live field rather than
    flattened to plain text, provided its exact visible text (e.g.
    "(Bhambra, 2021)") still appears somewhere in `text` - which is the
    normal case, since citation text itself should never be rephrased during
    translation, only the prose around it. Returns a list of citation fields
    that COULDN'T be matched (and were therefore flattened to plain text
    after all) - callers should surface this as a warning rather than fail
    silently, since it usually means the translation reworded a citation
    that should have been left untouched.
    """
    base_rpr = None
    if paragraph.runs:
        r = paragraph.runs[0]._r
        rpr = r.find(qn('w:rPr'))
        if rpr is not None:
            base_rpr = etree.tostring(rpr)

    fields = list(iter_field_spans(paragraph))  # [(visible_text, run_elements), ...]

    for r in list(paragraph.runs):
        r._r.getparent().remove(r._r)

    # Greedily tokenize `text` into plain-text chunks and field placeholders,
    # matching each field at most once, left to right, so a repeated citation
    # doesn't get spliced into the wrong occurrence.
    tokens = []  # ('text', str) | ('field', run_elements)
    used = [False] * len(fields)
    remaining = text
    while remaining:
        best_pos = best_idx = None
        for idx, (visible, _els) in enumerate(fields):
            if used[idx]:
                continue
            pos = remaining.find(visible)
            if pos != -1 and (best_pos is None or pos < best_pos):
                best_pos, best_idx = pos, idx
        if best_pos is None:
            tokens.append(('text', remaining))
            break
        if best_pos > 0:
            tokens.append(('text', remaining[:best_pos]))
        tokens.append(('field', fields[best_idx][1]))
        used[best_idx] = True
        remaining = remaining[best_pos + len(fields[best_idx][0]):]

    for kind, payload in tokens:
        if kind == 'field':
            for el in payload:
                paragraph._p.append(copy.deepcopy(el))
            continue
        # Split on [[highlight]] first, then *italic* within each resulting
        # piece - see the module-level note on why these don't nest.
        hl_pieces = _HIGHLIGHT_RE.split(payload)
        for h_i, hl_piece in enumerate(hl_pieces):
            if hl_piece == "":
                continue
            is_highlighted = (h_i % 2 == 1)
            pieces = _MARKUP_RE.split(hl_piece)  # alternating: plain, italic, plain, italic, ...
            for i, piece in enumerate(pieces):
                if piece == "":
                    continue
                run = paragraph.add_run(piece)
                if base_rpr is not None:
                    run._r.insert(0, etree.fromstring(base_rpr))
                if i % 2 == 1:  # odd index = came from inside *...*
                    run.italic = True
                if is_highlighted:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    return [visible for idx, (visible, _els) in enumerate(fields) if not used[idx]]


# --- Footnotes (word/footnotes.xml is not modeled by python-docx) -------------

_FOOTNOTES_PART = "word/footnotes.xml"


def extract_footnotes(docx_path):
    """Return [(footnote_id:str, text:str), ...] for real footnotes
    (skips separator/continuation placeholders ids 0/1)."""
    results = []
    with zipfile.ZipFile(docx_path) as z:
        if _FOOTNOTES_PART not in z.namelist():
            return results
        xml = z.read(_FOOTNOTES_PART)
    root = etree.fromstring(xml)
    for fn in root.findall(f'{W_NS}footnote'):
        fid = fn.get(f'{W_NS}id')
        if fid in ("0", "1", "-1"):
            continue
        texts = [t.text or "" for t in fn.iter(f'{W_NS}t')]
        results.append((fid, "".join(texts).strip()))
    return results


def write_footnotes(src_path, dst_path, id_to_text):
    """Copy src_path to dst_path, replacing footnote body text for each
    footnote id present in id_to_text (a dict of footnote_id -> new text,
    may contain *italic* markup). Runs within a footnote paragraph are
    collapsed to a single run (formatting of the original first run is kept),
    same approach as apply_translated_text."""
    if not id_to_text:
        if src_path != dst_path:
            shutil.copyfile(src_path, dst_path)
        return

    with zipfile.ZipFile(src_path) as zin:
        names = zin.namelist()
        if _FOOTNOTES_PART not in names:
            if src_path != dst_path:
                shutil.copyfile(src_path, dst_path)
            return
        data = {n: zin.read(n) for n in names}

    root = etree.fromstring(data[_FOOTNOTES_PART])
    for fn in root.findall(f'{W_NS}footnote'):
        fid = fn.get(f'{W_NS}id')
        if fid not in id_to_text:
            continue
        paras = fn.findall(f'{W_NS}p')
        if not paras:
            continue
        target_p = paras[0]
        runs = target_p.findall(f'{W_NS}r')
        base_rpr = None
        if runs:
            rpr = runs[0].find(f'{W_NS}rPr')
            if rpr is not None:
                base_rpr = etree.tostring(rpr)
        for r in runs:
            target_p.remove(r)

        pieces = _MARKUP_RE.split(id_to_text[fid])
        for i, piece in enumerate(pieces):
            if piece == "":
                continue
            new_r = etree.SubElement(target_p, f'{W_NS}r')
            if base_rpr is not None:
                rpr_el = etree.fromstring(base_rpr)
                if i % 2 == 1:
                    it = rpr_el.find(f'{W_NS}i')
                    if it is None:
                        etree.SubElement(rpr_el, f'{W_NS}i')
                new_r.append(rpr_el)
            elif i % 2 == 1:
                rpr_el = etree.SubElement(new_r, f'{W_NS}rPr')
                etree.SubElement(rpr_el, f'{W_NS}i')
            t_el = etree.SubElement(new_r, f'{W_NS}t')
            t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t_el.text = piece

    data[_FOOTNOTES_PART] = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)

    with zipfile.ZipFile(dst_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, content in data.items():
            zout.writestr(name, content)
